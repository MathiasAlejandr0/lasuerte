"""Genera el informe formal Word del proyecto Suertu2s (formato profesional + lenguaje claro)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "Informe-Proyecto-Suertu2s.docx"

FONT_TITLE = "Georgia"
FONT_BODY = "Calibri"
COLOR_GREEN = RGBColor(0x0E, 0x5C, 0x2A)
COLOR_GOLD = RGBColor(0xB8, 0x78, 0x17)
COLOR_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_MUTED = RGBColor(0x4A, 0x4A, 0x4A)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_ROW_ALT = "F4F7F4"
COLOR_HEADER_BG = "0E5C2A"
COLOR_PLAIN_BG = "F7F3E8"


def set_font(run, name, *, bold=False, italic=False, size=11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in (
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_TEXT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    pf = normal.paragraph_format
    pf.space_after = Pt(10)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for i, size, before, after in (
        (1, 18, 22, 12),
        (2, 13, 16, 8),
        (3, 12, 12, 6),
    ):
        style = styles[f"Heading {i}"]
        style.font.name = FONT_TITLE
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = COLOR_GREEN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_TITLE)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_TITLE)
        sp = style.paragraph_format
        sp.space_before = Pt(before)
        sp.space_after = Pt(after)
        sp.line_spacing = 1.15
        sp.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        if style_name in styles:
            st = styles[style_name]
            st.font.name = FONT_BODY
            st.font.size = Pt(11)
            st.paragraph_format.space_after = Pt(6)
            st.paragraph_format.space_before = Pt(2)
            st.paragraph_format.line_spacing = 1.15


def add_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = hp.add_run("SUERTU2S")
    set_font(r1, FONT_TITLE, bold=True, size=9, color=COLOR_GREEN)
    r2 = hp.add_run("  ·  Informe de proyecto (versión para todo el equipo)")
    set_font(r2, FONT_BODY, size=9, color=COLOR_MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "Documento confidencial — lenguaje técnico + explicaciones claras  ·  Página "
    )
    set_font(fr, FONT_BODY, size=8, color=COLOR_MUTED)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    set_font(run, FONT_BODY, size=8, color=COLOR_MUTED)
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_p(
    doc,
    text,
    *,
    bold=False,
    italic=False,
    size=11,
    align="left",
    space_after=10,
    space_before=0,
    color=None,
    font=None,
):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    run = p.add_run(text)
    set_font(
        run,
        font or FONT_BODY,
        bold=bold,
        italic=italic,
        size=size,
        color=color or COLOR_TEXT,
    )
    return p


def add_plain(doc, text):
    """Caja de explicación en lenguaje cotidiano para staff no técnico."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade_cell(cell, COLOR_PLAIN_BG)
    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    label = p.add_run("En palabras simples: ")
    set_font(label, FONT_TITLE, bold=True, size=10.5, color=COLOR_GREEN)
    body = p.add_run(text)
    set_font(body, FONT_BODY, size=10.5, color=COLOR_TEXT)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(12)
    return table


def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B87817")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_font(run, FONT_BODY, size=11, color=COLOR_TEXT)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            set_font(run, FONT_BODY, size=11, color=COLOR_TEXT)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        shade_cell(cell, COLOR_HEADER_BG)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        set_font(run, FONT_TITLE, bold=True, size=10, color=COLOR_WHITE)

    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            set_cell_margins(cell)
            if r_i % 2 == 1:
                shade_cell(cell, COLOR_ROW_ALT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(val))
            set_font(run, FONT_BODY, size=10, color=COLOR_TEXT)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(12)
    return table


def build():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)
    add_header_footer(section)

    # Portada
    for _ in range(2):
        doc.add_paragraph()

    add_p(
        doc,
        "PROYECTO INFORMÁTICO",
        bold=True,
        size=12,
        align="center",
        color=COLOR_GOLD,
        font=FONT_TITLE,
        space_after=6,
    )
    add_rule(doc)
    add_p(
        doc,
        "SUERTU2S",
        bold=True,
        size=32,
        align="center",
        color=COLOR_GREEN,
        font=FONT_TITLE,
        space_after=8,
        space_before=8,
    )
    add_p(
        doc,
        "Plataforma de venta de ilustraciones digitales y sorteo",
        size=13,
        align="center",
        color=COLOR_MUTED,
        font=FONT_TITLE,
        italic=True,
        space_after=16,
    )
    add_p(
        doc,
        "Informe técnico-funcional · Costos · Riesgos · Seguridad · Mercado\n"
        "Solicitud formal de accesos (datos, hosting y DNS)",
        size=11,
        align="center",
        color=COLOR_TEXT,
        space_after=10,
    )
    add_p(
        doc,
        "Incluye explicaciones en lenguaje cotidiano para que todo el staff\n"
        "entienda el proyecto (no solo el área informática).",
        size=10.5,
        align="center",
        italic=True,
        color=COLOR_GREEN,
        space_after=22,
    )
    add_rule(doc)

    meses = (
        "enero febrero marzo abril mayo junio "
        "julio agosto septiembre octubre noviembre diciembre"
    ).split()
    hoy = date.today()
    meta = [
        f"Fecha: {hoy.day} de {meses[hoy.month - 1]} de {hoy.year}",
        "Versión: 1.1 — con lenguaje claro para el equipo",
        "Estado: Documento formal para presentación / trabajo",
        "Repositorio: https://github.com/MathiasAlejandr0/lasuerte",
    ]
    for line in meta:
        add_p(doc, line, align="center", size=10.5, color=COLOR_MUTED, space_after=4)

    add_p(doc, "", space_after=14)
    add_p(
        doc,
        "Confidencial — uso interno / entrega académica o laboral",
        italic=True,
        size=9,
        align="center",
        color=COLOR_GOLD,
        font=FONT_TITLE,
    )
    doc.add_page_break()

    # Cómo leer
    add_heading(doc, "Cómo leer este documento", 1)
    add_p(
        doc,
        "Este informe está pensado para un equipo mixto: hay términos informáticos "
        "(porque el proyecto lo requiere) y, en cada sección importante, un recuadro "
        "beige titulado “En palabras simples” que traduce la idea sin jerga.",
        align="justify",
    )
    add_plain(
        doc,
        "Si no eres de informática, puedes leer primero los recuadros beige y las "
        "conclusiones. Ahí está lo esencial: qué hace la web, cuánto cuesta referencialmente, "
        "qué riesgos hay y qué accesos necesitamos del resto del equipo.",
    )
    add_p(
        doc,
        "Al final hay un glosario amplio (Anexo A) con las palabras técnicas explicadas "
        "una a una.",
        align="justify",
    )

    # Índice
    add_heading(doc, "Índice de contenidos", 1)
    toc = [
        "1. Resumen ejecutivo",
        "2. Antecedentes y justificación",
        "3. Objetivos del proyecto",
        "4. Alcance",
        "5. Descripción funcional del sistema",
        "6. Arquitectura y stack tecnológico",
        "7. Comparación con WordPress / WooCommerce",
        "8. Análisis de mercado",
        "9. Análisis de costos",
        "10. Análisis de riesgos",
        "11. Seguridad de la información",
        "12. Plan de implementación y go-live",
        "13. Solicitud formal de accesos (datos, hosting, DNS y credenciales)",
        "14. Entregables",
        "15. Conclusiones y recomendaciones",
        "16. Anexos (incluye glosario para todo el staff)",
    ]
    for item in toc:
        add_p(doc, item, space_after=5, size=11, color=COLOR_TEXT)
    doc.add_page_break()

    # 1
    add_heading(doc, "1. Resumen ejecutivo", 1)
    add_p(
        doc,
        "Suertu2s es una plataforma web orientada a la venta de packs de ilustraciones "
        "digitales del sur de Chile, que además entrega números de participación (boletos) "
        "para un sorteo de premios. El proyecto corresponde a la migración y modernización "
        "del sitio original basado en WordPress/WooCommerce hacia una aplicación a medida "
        "desarrollada con Next.js 16, TypeScript, Tailwind CSS y Supabase.",
        align="justify",
    )
    add_plain(
        doc,
        "Hoy tenemos una página propia (no WordPress) donde la gente compra una "
        "ilustración digital y, con esa compra, recibe números para un sorteo. También "
        "hay un panel para administrar pedidos, afiliados y el sorteo, y un espacio para "
        "vendedores/embajadores. El código ya está avanzado; para publicarlo “de verdad” "
        "faltan accesos y claves que debe entregar el cliente o el área a cargo del dominio "
        "y las cuentas de pago.",
    )
    add_p(
        doc,
        "El sistema contempla checkout con Mercado Pago y Webpay, asignación de números, "
        "consulta de tickets, emails de confirmación, panel de administración operativo, "
        "portal de afiliados, transmisión en vivo del sorteo, anuncio de ganador y "
        "controles de seguridad alineados a un producto de e-commerce/sorteo.",
        align="justify",
    )
    add_p(
        doc,
        "Valor orientativo de desarrollo en Chile: $3.200.000 a $7.000.000 CLP (+ IVA), "
        "según perfil del proveedor (freelancer senior, estudio o agencia). El go-live "
        "requiere credenciales y accesos que se solicitan formalmente en este documento.",
        align="justify",
    )
    add_plain(
        doc,
        "“Go-live” significa poner la web en internet con el dominio real y cobros reales. "
        "El rango de plata es una referencia de mercado de cuánto costaría encargar un "
        "proyecto así en Chile; no es una factura automática.",
    )

    # 2
    add_heading(doc, "2. Antecedentes y justificación", 1)
    add_heading(doc, "2.1 Situación anterior", 2)
    add_bullets(
        doc,
        [
            "Sitio público en WordPress / WooCommerce (suertu2s.cl).",
            "Dependencia de tema y plugins para carrito, pagos y funcionalidades.",
            "Dificultad para modelar reglas propias de sorteo, tickets y afiliados.",
            "Mayor superficie de ataque y dependencia de actualizaciones de terceros.",
        ],
    )
    add_plain(
        doc,
        "Antes la web funcionaba sobre WordPress: un sistema genérico de páginas al que "
        "se le agregan “plugins” (complementos). Eso sirve para tiendas simples, pero "
        "para un sorteo con números, afiliados y reglas propias se vuelve frágil: si un "
        "plugin se actualiza mal o tiene un agujero de seguridad, puede afectar todo.",
    )
    add_heading(doc, "2.2 Situación propuesta", 2)
    add_bullets(
        doc,
        [
            "Aplicación a medida con lógica de negocio en servidor.",
            "Integraciones de pago propias (Mercado Pago y Webpay).",
            "Panel admin y portal de afiliados diseñados para operar el sorteo.",
            "Base de datos Postgres (Supabase) con RLS y migración única de esquema.",
            "Mejor trazabilidad, seguridad y capacidad de evolución del producto.",
        ],
    )
    add_plain(
        doc,
        "Ahora la web está hecha a la medida de Suertu2s: las reglas importantes "
        "(precios, números del sorteo, quién puede administrar) viven en nuestro código "
        "y en una base de datos propia, no en una colección de plugins ajenos.",
    )
    add_heading(doc, "2.3 Justificación", 2)
    add_p(
        doc,
        "Para un negocio de sorteos con packs, números, afiliados y transmisión en vivo, "
        "una solución ensamblada con plugins es frágil y difícil de auditar. Una app a "
        "medida reduce riesgo operativo, mejora control de pagos/tickets y proyecta una "
        "imagen más profesional frente a usuarios y stakeholders.",
        align="justify",
    )
    add_plain(
        doc,
        "En corto: no elegimos la tecnología “porque suena moderna”, sino porque el "
        "negocio necesita control, seguridad y operación clara. Eso se traduce en menos "
        "sorpresas y una imagen más seria frente a clientes y socios.",
    )

    # 3
    add_heading(doc, "3. Objetivos del proyecto", 1)
    add_heading(doc, "3.1 Objetivo general", 2)
    add_p(
        doc,
        "Diseñar, desarrollar e implementar una plataforma web profesional que permita "
        "vender packs digitales, gestionar participación en sorteo, cobrar con pasarelas "
        "chilenas/regionales y operar el negocio desde un panel administrativo seguro.",
        align="justify",
    )
    add_plain(
        doc,
        "El objetivo es tener una web que venda, cobre bien, entregue números del sorteo, "
        "avise por correo y permita al equipo administrar todo sin depender de un técnico "
        "para cada tarea del día a día.",
    )
    add_heading(doc, "3.2 Objetivos específicos", 2)
    add_numbered(
        doc,
        [
            "Migrar la experiencia comercial desde WordPress a Next.js + Supabase.",
            "Implementar checkout con selector real de Mercado Pago y Webpay.",
            "Asignar y consultar números de sorteo de forma confiable.",
            "Enviar emails de confirmación con reintento/flag de entrega.",
            "Entregar panel admin (pedidos, analítica, afiliados, ajustes, live, ganador).",
            "Endurecer seguridad (sesiones, CSRF, CSP, webhooks firmados, RLS).",
            "Documentar costos, riesgos, mercado y solicitud formal de accesos.",
        ],
    )
    add_plain(
        doc,
        "Traducción: sacar la tienda de WordPress; que el cliente elija cómo pagar "
        "(Mercado Pago o Webpay); que los números del sorteo se asignen bien; que el "
        "correo de confirmación se pueda reintentar si falla; que exista un panel de "
        "control; que la web esté protegida; y que este documento deje todo claro al equipo.",
    )

    # 4
    add_heading(doc, "4. Alcance", 1)
    add_plain(
        doc,
        "“Alcance” significa qué entra en este proyecto y qué no. Evita malentendidos "
        "del tipo “pensé que también incluía X”.",
    )
    add_heading(doc, "4.1 Incluido", 2)
    add_bullets(
        doc,
        [
            "Landing, carrito, checkout, consulta de números, páginas legales/sorteo.",
            "Pagos (MP/Webpay), webhooks, mock seguro solo en desarrollo.",
            "Admin: resumen, analítica, pedidos, clientes, tickets, afiliados, ajustes.",
            "Portal afiliados con login, comisiones y QR.",
            "Live del sorteo, cierre, anuncio de ganador y bloqueo de compras al cerrar.",
            "Migración SQL completa, CI (typecheck/lint/tests), documentación.",
        ],
    )
    add_heading(doc, "4.2 Excluido / pendiente de terceros", 2)
    add_bullets(
        doc,
        [
            "Provisión de cuentas y claves de Mercado Pago, Webpay, Supabase y Resend.",
            "Gestión DNS del dominio suertu2s.cl y certificados asociados al hosting.",
            "Revisión legal definitiva de bases del sorteo por abogado.",
            "2FA admin, rate-limit distribuido (Redis) y SEO avanzado (fase posterior).",
            "Persistencia total del catálogo editable admin en Supabase (mejora futura).",
        ],
    )
    add_plain(
        doc,
        "Lo que ya está hecho es el “cerebro” de la web. Lo que pide el cliente/TI son "
        "las llaves: cuentas de pago, base de datos en la nube, correo, dominio y DNS. "
        "También queda fuera la revisión legal de un abogado y algunas mejoras futuras "
        "(como doble verificación al entrar al admin).",
    )

    # 5
    add_heading(doc, "5. Descripción funcional del sistema", 1)
    add_plain(
        doc,
        "Esta sección cuenta qué puede hacer cada tipo de persona en la plataforma "
        "y cómo es el camino de una compra, sin entrar todavía en detalles técnicos profundos.",
    )
    add_heading(doc, "5.1 Actores", 2)
    add_table(
        doc,
        ["Actor", "Rol principal", "En simple"],
        [
            [
                "Comprador",
                "Compra packs, recibe números y emails, consulta tickets",
                "La persona que paga en la web",
            ],
            [
                "Administrador",
                "Opera pedidos, sorteo, afiliados, live y ganador",
                "Quien maneja el negocio desde /admin",
            ],
            [
                "Afiliado",
                "Promociona con código, ve ventas y comisiones",
                "Embajador / vendedor con link propio",
            ],
            [
                "Sistema / pasarelas",
                "Confirma pagos vía webhook o retorno Webpay",
                "Mercado Pago o Webpay avisando “este pago quedó OK”",
            ],
        ],
        col_widths=[3.2, 6.5, 6.3],
    )

    add_heading(doc, "5.2 Flujo de compra", 2)
    add_numbered(
        doc,
        [
            "El usuario elige un pack (ilustración + N números).",
            "Revisa carrito y entra a checkout.",
            "Acepta privacidad, bases y mayoría de edad.",
            "Selecciona Mercado Pago o Webpay.",
            "Paga; el backend confirma, asigna números y envía email.",
            "Puede consultar números con su correo en /check-tickets.",
        ],
    )
    add_plain(
        doc,
        "Es como una tienda online normal, con un plus: al pagar no solo “compra el "
        "dibujo”, también recibe números para el sorteo y un correo con esa información. "
        "Después puede volver a la web, poner su correo y ver sus números.",
    )

    add_heading(doc, "5.3 Módulos principales", 2)
    add_table(
        doc,
        ["Módulo", "Funciones", "Para el equipo significa…"],
        [
            [
                "Sitio público",
                "Landing, packs, FAQ, countdown, live, ganador, widget consulta, FAB soporte",
                "Lo que ve cualquier visitante",
            ],
            [
                "Checkout",
                "Carrito, datos, selector de pago, validaciones",
                "La caja para pagar",
            ],
            [
                "Tickets",
                "Asignación atómica, consulta por email (página y widget)",
                "Entrega y consulta de números del sorteo",
            ],
            [
                "Emails",
                "Confirmación con ilustraciones + reintento",
                "Correo automático después del pago",
            ],
            [
                "Admin",
                "KPIs, pedidos, analítica, sorteos, ajustes, afiliados",
                "Escritorio de control del negocio",
            ],
            [
                "Afiliados",
                "Portal, comisiones, liquidaciones, QR",
                "Herramienta para embajadores",
            ],
            [
                "Sorteo",
                "Ciclos, historial, cierre, ganador, bloqueo de compras, live",
                "Control del ciclo activo y la transmisión",
            ],
        ],
        col_widths=[3, 6.5, 6.5],
    )

    add_heading(doc, "5.4 Experiencia de landing (integraciones recientes v2)", 2)
    add_plain(
        doc,
        "Además del flujo de compra, la versión 2 incorpora piezas de experiencia pensadas "
        "para que el visitante consulte su pedido y contacte soporte sin salir de la home, "
        "y para que el equipo pueda mostrar una demo más completa.",
    )
    add_table(
        doc,
        ["Pieza", "Qué hace", "Para el equipo significa…"],
        [
            [
                "Título del hero",
                "“Compra Tus Ilustraciones Digitales Y Participa En El Sorteo De Nuestra Moto”",
                "Mensaje comercial claro en la primera pantalla",
            ],
            [
                "Widget Consulta tu pedido",
                "Email → busca códigos; FAQ rápida; botón WhatsApp",
                "Menos fricción para quien ya compró",
            ],
            [
                "Botón flotante (FAB)",
                "Burbuja de soporte abajo a la derecha; abre el widget",
                "Acceso permanente a consulta/contacto",
            ],
            [
                "Carrusel hero",
                "Flechas con hover dorado (marca)",
                "Detalle visual alineado a la identidad",
            ],
            [
                "Admin → Sorteos",
                "Ciclo activo, nuevo ciclo e historial",
                "Operar varios sorteos sin mezclar pedidos",
            ],
            [
                "Repo v1 / v2",
                "Carpetas independientes para comparar versiones",
                "Elegir con cuál quedarse en reunión",
            ],
        ],
        col_widths=[3.5, 6.5, 6],
    )
    add_plain(
        doc,
        "WhatsApp del widget: opcional vía NEXT_PUBLIC_WHATSAPP_NUMBER (dígitos con código de país). "
        "Sin número configurado el botón igual abre WhatsApp con un mensaje prearmado.",
    )

    add_heading(doc, "5.5 Rutas relevantes", 2)
    add_plain(
        doc,
        "Una “ruta” es la dirección que se escribe después del dominio. "
        "Ejemplo: suertu2s.cl/checkout es la página de pago.",
    )
    add_table(
        doc,
        ["Ruta", "Descripción", "Quién la usa"],
        [
            ["/", "Landing (página de inicio)", "Público"],
            ["/carrito", "Carrito", "Comprador"],
            ["/checkout", "Pago", "Comprador"],
            ["/check-tickets", "Consulta de números", "Comprador"],
            ["/afiliados", "Portal afiliado", "Embajador"],
            ["/admin", "Panel administrador (incluye Sorteos)", "Equipo interno"],
            ["/pago/exito o /pago/error", "Resultado de pago", "Comprador"],
            ["/#consulta-codigos", "Widget consulta en la home", "Comprador"],
        ],
        col_widths=[5, 6, 5],
    )

    # 6
    add_heading(doc, "6. Arquitectura y stack tecnológico", 1)
    add_plain(
        doc,
        "“Stack” = el conjunto de herramientas con las que está construida la web. "
        "No hace falta memorizar los nombres: lo importante es que son tecnologías "
        "actuales, usadas por muchas empresas, y permiten cobros, base de datos y "
        "seguridad de nivel profesional.",
    )
    add_p(
        doc,
        "Arquitectura: aplicación full-stack Next.js (App Router) desplegable en "
        "hosting Node/Vercel, con API routes, middleware de seguridad, base Postgres "
        "en Supabase (service role en servidor) y proveedores externos de pago/email.",
        align="justify",
    )
    add_table(
        doc,
        ["Capa", "Tecnología", "En simple"],
        [
            [
                "Frontend/Backend",
                "Next.js 16, React, TypeScript, Tailwind CSS 4",
                "El programa de la web (pantallas + lógica)",
            ],
            [
                "Datos",
                "Supabase (Postgres + RLS) / modo memoria en demo",
                "La “bodega” donde se guardan pedidos y números",
            ],
            [
                "Pagos",
                "Mercado Pago Checkout Pro, Transbank Webpay Plus",
                "Las cajas de cobro que ya conoce la gente",
            ],
            [
                "Email",
                "Resend",
                "Servicio que envía los correos automáticos",
            ],
            [
                "Estado cliente",
                "Zustand (carrito)",
                "Recuerda qué packs agregaste antes de pagar",
            ],
            [
                "Calidad",
                "ESLint, Prettier, Vitest, Playwright, GitHub Actions",
                "Revisiones automáticas para detectar errores",
            ],
            [
                "Código fuente",
                "GitHub: MathiasAlejandr0/lasuerte",
                "Dónde está guardado el proyecto en internet",
            ],
        ],
        col_widths=[3.2, 6.8, 6],
    )

    # 7
    add_heading(doc, "7. Comparación con WordPress / WooCommerce", 1)
    add_plain(
        doc,
        "WordPress es más barato y rápido de armar para una tienda simple. Suertu2s "
        "cuesta más porque es un sistema hecho a medida para sorteo, números, afiliados "
        "y operación. No es “lujo”: es control y profesionalismo para este negocio.",
    )
    add_table(
        doc,
        ["Dimensión", "WordPress + Woo", "Suertu2s (Next.js + Supabase)"],
        [
            ["Naturaleza", "CMS + plugins", "App a medida para sorteo/packs"],
            [
                "Lógica crítica",
                "Dispersa en plugins",
                "Servidor controlado (precios, tickets)",
            ],
            [
                "Pagos Chile",
                "Calidad variable por plugin",
                "Integración propia + validación de monto",
            ],
            [
                "Seguridad",
                "Alta dependencia de updates",
                "Sesiones, CSRF, CSP, RLS, webhooks",
            ],
            [
                "Operación sorteo",
                "Improvisada",
                "Live, cierre, ganador, bloqueo ventas",
            ],
            [
                "Costo inicial típico",
                "$0,8M – $2,5M CLP",
                "$3,2M – $7M CLP",
            ],
            [
                "Mantenibilidad",
                "Riesgo por updates WP/plugins",
                "Evolución controlada en código",
            ],
        ],
        col_widths=[3.5, 5.5, 7],
    )
    add_p(
        doc,
        "Conclusión comparativa: WordPress es más barato para tiendas simples; para "
        "Suertu2s, la opción a medida es más profesional, auditable y alineada al negocio.",
        align="justify",
    )

    # 8
    add_heading(doc, "8. Análisis de mercado", 1)
    add_plain(
        doc,
        "Aquí no hablamos de código, sino del negocio: a quién le sirve esto, por qué "
        "tiene sentido y en qué se diferencia de otras opciones.",
    )
    add_heading(doc, "8.1 Problema / oportunidad", 2)
    add_p(
        doc,
        "Existe demanda local por experiencias de compra digital con incentivo de "
        "sorteo (packs + participación). El mercado chileno exige medios de pago "
        "conocidos (Webpay/Mercado Pago), confianza visual y transparencia del proceso "
        "de sorteo (live / anuncio de ganador).",
        align="justify",
    )
    add_heading(doc, "8.2 Público objetivo", 2)
    add_bullets(
        doc,
        [
            "Compradores digitales 18+ interesados en ilustración y/o sorteos.",
            "Audiencia regional (sur de Chile) y nacional por canal digital.",
            "Embajadores/afiliados que promueven con código de referido.",
            "Operadores del negocio que necesitan panel y métricas claras.",
        ],
    )
    add_heading(doc, "8.3 Propuesta de valor", 2)
    add_bullets(
        doc,
        [
            "Producto digital inmediato (ilustración) + chance de premio.",
            "Pagos locales confiables.",
            "Transparencia del sorteo (countdown, live, ganador).",
            "Canal de afiliados medible.",
        ],
    )
    add_heading(doc, "8.4 Competencia / alternativas", 2)
    add_bullets(
        doc,
        [
            "Tiendas WooCommerce genéricas con sorteo manual o plugin.",
            "Landing + formularios sin asignación robusta de tickets.",
            "Plataformas extranjeras poco adaptadas a Webpay/CLP.",
        ],
    )
    add_heading(doc, "8.5 Diferenciación", 2)
    add_p(
        doc,
        "Suertu2s concentra venta, tickets, afiliados, operación del sorteo y "
        "seguridad en un solo producto, con experiencia de marca propia (no plantilla).",
        align="justify",
    )
    add_plain(
        doc,
        "La diferencia frente a una tienda genérica es que aquí el sorteo no es un "
        "parche: está integrado (números, live, cierre, ganador, afiliados y panel).",
    )

    # 9
    add_heading(doc, "9. Análisis de costos", 1)
    add_plain(
        doc,
        "Hay dos tipos de plata: (1) lo que cuesta construir el sistema una vez, y "
        "(2) lo que cuesta mantenerlo mes a mes (hosting, base de datos, correos, "
        "comisiones de las pasarelas). Los montos son referenciales de mercado en Chile.",
    )
    add_heading(doc, "9.1 Inversión de desarrollo (pago único, + IVA)", 2)
    add_table(
        doc,
        ["Enfoque", "Rango CLP", "Observación"],
        [
            ["WordPress básico", "$800.000 – $2.500.000", "Tema + Woo; lógica limitada"],
            [
                "WordPress a medida",
                "$2.000.000 – $4.500.000",
                "Plugins custom / PHP",
            ],
            [
                "Next.js + Supabase (Suertu2s)",
                "$3.200.000 – $7.000.000",
                "App completa actual",
            ],
        ],
        col_widths=[5, 5, 6],
    )
    add_heading(doc, "9.2 Costos operativos mensuales estimados", 2)
    add_table(
        doc,
        ["Ítem", "Estimación mensual", "Notas"],
        [
            ["Hosting app (Vercel/similar)", "USD 0 – 20+", "Donde “vive” la web"],
            ["Supabase", "USD 0 – 25+", "Base de datos en la nube"],
            ["Dominio + DNS", "CLP bajo / anual", "El nombre suertu2s.cl"],
            ["Resend / email", "USD 0 – 20+", "Envío de correos"],
            ["Pasarelas", "% comisión por venta", "Lo cobra MP/Transbank por cobro"],
            ["Mantención técnica", "acordar", "Soporte y mejoras"],
        ],
        col_widths=[5, 4.5, 6.5],
    )
    add_heading(doc, "9.3 Beneficios económicos esperados del stack", 2)
    add_bullets(
        doc,
        [
            "Menor costo de incidentes por plugins rotos o vulnerabilidades comunes.",
            "Mayor conversión potencial por checkout claro y confianza de pago.",
            "Operación más eficiente (admin + afiliados + analítica).",
            "Base reutilizable para nuevas campañas/sorteos.",
        ],
    )
    add_plain(
        doc,
        "Pagar más al inicio puede salir más barato después: menos caídas, menos "
        "emergencias y menos tiempo perdido parcheando plugins. Además, un checkout "
        "claro ayuda a que más gente termine la compra.",
    )

    # 10
    add_heading(doc, "10. Análisis de riesgos", 1)
    add_plain(
        doc,
        "Un riesgo no es una predicción de fracaso: es algo que podría pasar y conviene "
        "tener plan B. “Impacto” = qué tan grave sería. “Prob.” = qué tan posible se ve.",
    )
    add_table(
        doc,
        ["Riesgo", "Impacto", "Prob.", "Mitigación (qué hacemos)"],
        [
            [
                "Falta de claves de pago/Supabase a tiempo",
                "Alto",
                "Media",
                "Solicitud formal (sección 13) y checklist go-live",
            ],
            [
                "Webhook/email fallido",
                "Alto",
                "Media",
                "Firma MP, flag email, reintento, reenvío admin",
            ],
            [
                "Fraude o montos alterados",
                "Alto",
                "Baja",
                "Precios en servidor + validación de monto",
            ],
            [
                "Acceso indebido al admin",
                "Alto",
                "Baja",
                "Sesión firmada, lista de emails, anti-CSRF, secretos fuertes",
            ],
            [
                "Sorteo cerrado pero siguen ventas",
                "Medio",
                "Baja",
                "La web y el servidor bloquean compras si está cerrado",
            ],
            [
                "Catálogo admin en archivo local (serverless)",
                "Medio",
                "Media",
                "Mejora futura: guardar catálogo en Supabase",
            ],
            [
                "Incumplimiento legal de bases",
                "Alto",
                "Media",
                "Revisión con abogado",
            ],
            [
                "Caída de hosting/DNS",
                "Alto",
                "Baja",
                "Proveedor confiable, DNS bien configurado, monitoreo",
            ],
        ],
        col_widths=[4.5, 2, 2, 7.5],
    )
    add_plain(
        doc,
        "El riesgo más práctico hoy no es técnico: es que falten a tiempo las claves "
        "y accesos (pagos, base de datos, dominio). Por eso existe la sección 13.",
    )

    # 11
    add_heading(doc, "11. Seguridad de la información", 1)
    add_plain(
        doc,
        "Seguridad, en este proyecto, significa: que no cualquiera entre al panel, "
        "que no se pueda falsear un pago, que los datos de clientes no queden expuestos "
        "y que las contraseñas/claves no anden sueltas en chats o en GitHub.",
    )
    add_heading(doc, "11.1 Controles implementados", 2)
    add_bullets(
        doc,
        [
            "Autenticación admin con cookie httpOnly, SameSite=Lax y firma HMAC.",
            "ADMIN_SESSION_SECRET obligatorio en producción (separado de la password).",
            "Soporte de ADMIN_PASSWORD_HASH (scrypt).",
            "Middleware: protección de APIs admin y chequeo same-origin (anti-CSRF).",
            "Mock de pagos desactivado en producción; token firmado en desarrollo.",
            "Webhook Mercado Pago con verificación de firma y monto.",
            "RLS en Supabase; assign_tickets solo service_role.",
            "Passwords de afiliados con scrypt.",
            "Headers: CSP, HSTS, X-Frame-Options, nosniff, etc.",
            "Rate limiting en login, checkout y webhooks.",
            "Secretos fuera de git (.env ignorado; solo .env.example).",
        ],
    )
    add_plain(
        doc,
        "Traducción de lo más importante: el panel admin pide usuario/clave y deja "
        "una “sesión” protegida; en internet real no se puede simular un pago de prueba; "
        "Mercado Pago debe firmar sus avisos; la base de datos no deja leer pedidos a "
        "cualquiera; y las claves secretas no se suben al repositorio público.",
    )
    add_heading(doc, "11.2 Datos personales tratados", 2)
    add_bullets(
        doc,
        [
            "Email, nombre, RUT/teléfono (según captura), datos de pedido y pago externo.",
            "Números de sorteo asociados al comprador.",
            "Datos de afiliados (email, comisiones, liquidaciones).",
        ],
    )
    add_plain(
        doc,
        "Guardamos lo necesario para vender, entregar números y operar afiliados. "
        "Eso implica responsabilidad: no compartir bases de clientes ni pegar claves "
        "en WhatsApp o correos abiertos.",
    )
    add_heading(doc, "11.3 Principios", 2)
    add_p(
        doc,
        "Mínimo privilegio, fail-closed sin configuración admin, no exponer secretos al "
        "cliente, validar siempre en servidor y registrar errores sin filtrar datos "
        "sensibles en producción.",
        align="justify",
    )
    add_plain(
        doc,
        "Principios en castellano: cada persona solo ve lo que necesita; si falta "
        "configuración de seguridad, el sistema prefiere negarse a abrir antes que "
        "quedar abierto; y los precios/números se deciden en el servidor, no en el "
        "navegador del cliente (para que nadie los manipule).",
    )

    # 12
    add_heading(doc, "12. Plan de implementación y go-live", 1)
    add_plain(
        doc,
        "Esto es el orden recomendado para llegar a producción sin improvisar. "
        "Cada fase tiene responsables distintos: desarrollo y negocio/TI del cliente.",
    )
    add_table(
        doc,
        ["Fase", "Actividades", "Responsable", "En simple"],
        [
            [
                "1. Desarrollo",
                "App, admin, pagos, seguridad, docs",
                "Equipo desarrollo",
                "Construir el sistema",
            ],
            [
                "2. Accesos",
                "DNS, hosting, Supabase, MP, Webpay, Resend",
                "Cliente / infra (§13)",
                "Entregar las llaves",
            ],
            [
                "3. Configuración",
                "Env prod, migración SQL, dominio, webhooks",
                "Dev + cliente",
                "Conectar todo",
            ],
            [
                "4. Pruebas",
                "Pago real controlado, emails, admin, afiliados",
                "Dev + stakeholders",
                "Probar antes del lanzamiento",
            ],
            [
                "5. Go-live",
                "DNS apuntando, mock off, monitoreo inicial",
                "Equipo conjunto",
                "Abrir al público",
            ],
            [
                "6. Operación",
                "Soporte, sorteo, live, cierre/ganador",
                "Operación negocio",
                "Usar el sistema día a día",
            ],
        ],
        col_widths=[2.4, 5.2, 3.4, 5],
    )

    # 13
    add_heading(
        doc,
        "13. Solicitud formal de accesos (datos, hosting, DNS y credenciales)",
        1,
    )
    add_plain(
        doc,
        "Esta sección es una carta formal al resto del equipo/cliente. No pide “todo "
        "el control de la empresa”: pide lo mínimo para publicar la web con cobros "
        "reales, dominio propio y correos funcionando. Las claves deben enviarse por "
        "un canal seguro, no por un chat grupal.",
    )
    add_p(
        doc,
        "Por medio de la presente, se solicita formalmente a la contraparte responsable "
        "(cliente, área de TI, o titular del dominio/marca Suertu2s) la entrega o "
        "habilitación de los accesos e información necesarios para configurar, desplegar "
        "y operar el proyecto en un entorno productivo seguro.",
        align="justify",
    )
    add_p(
        doc,
        "La solicitud se realiza con fines estrictamente técnicos y operativos del "
        "proyecto informático Suertu2s. Los accesos deben entregarse por canal seguro "
        "(gestor de secretos, correo cifrado o reunión controlada), evitando chats "
        "públicos o repositorios.",
        align="justify",
    )

    add_heading(doc, "13.1 Acceso a datos y contenidos", 2)
    add_bullets(
        doc,
        [
            "Textos legales vigentes (bases del sorteo, privacidad, términos).",
            "Imágenes/ilustraciones finales y derechos de uso.",
            "Datos de packs (nombres, precios, cantidad de números) si difieren del seed.",
            "Listado de emails admin autorizados.",
            "Información de afiliados iniciales (si aplica).",
            "Calendario del sorteo (fecha/hora de cierre) y criterios de anuncio.",
        ],
    )
    add_plain(
        doc,
        "Necesitamos los textos e imágenes finales, quiénes pueden entrar al panel, "
        "y la fecha oficial de cierre del sorteo. Sin eso, la web puede quedar con "
        "contenido provisional.",
    )

    add_heading(doc, "13.2 Hosting y despliegue", 2)
    add_bullets(
        doc,
        [
            "Cuenta de hosting/plataforma (p. ej. Vercel u equivalente Node) o invitación al proyecto.",
            "Permisos para crear variables de entorno de producción.",
            "Confirmación de región/plan y límites de tráfico esperados.",
            "Acceso a logs básicos para diagnóstico post-despliegue.",
        ],
    )
    add_plain(
        doc,
        "Hosting = el “arriendo” del servidor donde corre la web. Las variables de "
        "entorno son las casillas secretas donde se pegan las claves (pagos, base de "
        "datos, etc.) sin meterlas en el código público.",
    )

    add_heading(doc, "13.3 DNS y dominio", 2)
    add_bullets(
        doc,
        [
            "Acceso al panel DNS del dominio (p. ej. suertu2s.cl) o gestión por tercero autorizado.",
            "Permiso para crear/modificar registros A/CNAME/TXT necesarios (app, www, email).",
            "Coordinación de SSL/TLS (habitualmente automático en el hosting).",
            "Confirmación del dominio canónico y redirecciones www/apex.",
        ],
    )
    add_plain(
        doc,
        "El dominio es el nombre (suertu2s.cl). El DNS es la “agenda” que dice a qué "
        "servidor apunta ese nombre. SSL es el candadito https de seguridad en el navegador.",
    )

    add_heading(doc, "13.4 Credenciales de servicios", 2)
    add_table(
        doc,
        ["Servicio", "Qué pedimos", "Para qué"],
        [
            [
                "Supabase",
                "URL, claves y permiso para aplicar la migración SQL",
                "Guardar pedidos, números, afiliados",
            ],
            [
                "Mercado Pago",
                "Access token y secreto del webhook",
                "Cobrar y confirmar pagos MP",
            ],
            [
                "Webpay/Transbank",
                "Commerce code, API key y ambiente",
                "Cobrar con Webpay",
            ],
            [
                "Resend",
                "API key y correo remitente verificado",
                "Enviar emails de compra",
            ],
            [
                "Admin app",
                "Emails autorizados, clave y secreto de sesión",
                "Entrar al panel /admin",
            ],
        ],
        col_widths=[3.5, 7, 5.5],
    )

    add_heading(doc, "13.5 Variables de entorno mínimas de producción", 2)
    add_plain(
        doc,
        "La lista siguiente son nombres técnicos de esas “casillas secretas”. "
        "El informático las configura; el resto del equipo solo debe saber que existen "
        "y que no se comparten en público.",
    )
    add_bullets(
        doc,
        [
            "NEXT_PUBLIC_SITE_URL — dirección pública de la web",
            "Claves Supabase (URL, anon, service role)",
            "Claves Mercado Pago (token + webhook secret)",
            "Claves Webpay (commerce code, API key, env=production)",
            "Claves Resend (API key + EMAIL_FROM)",
            "Admin: ADMIN_EMAILS, password/hash, ADMIN_SESSION_SECRET",
            "TRUST_PROXY=true si hay CDN/proxy",
            "AFFILIATE_SESSION_SECRET (recomendado)",
        ],
    )

    add_heading(doc, "13.6 Compromisos del equipo solicitante", 2)
    add_bullets(
        doc,
        [
            "Usar los accesos solo para el objeto del proyecto.",
            "No versionar secretos en Git.",
            "Aplicar principio de mínimo privilegio.",
            "Devolver/rotar credenciales si deja de ser necesario el acceso.",
            "Informar incidentes de seguridad de forma inmediata.",
        ],
    )
    add_plain(
        doc,
        "Compromiso claro: las llaves no se suben a GitHub, no se reenvían a terceros "
        "sin necesidad, y si algo se filtra se avisa al tiro.",
    )

    add_heading(doc, "13.7 Formato de respuesta solicitado", 2)
    add_p(
        doc,
        "Se solicita respuesta escrita (correo o acta) indicando: (a) accesos otorgados, "
        "(b) responsable de cada servicio, (c) fecha de habilitación, (d) restricciones "
        "o condiciones de uso. Con esa información se podrá completar el checklist "
        "“Listo para operar” del panel admin y proceder al go-live.",
        align="justify",
    )
    add_plain(
        doc,
        "Pedimos una respuesta por escrito del estilo: “te dimos acceso a X el día Y, "
        "el responsable es Z”. Así queda registro y no se pierde en una conversación verbal.",
    )

    add_p(doc, "", space_after=6)
    add_p(doc, "Sin otro particular, se agradece la gestión.", align="justify")
    add_p(doc, "", space_after=16)
    add_p(doc, "_______________________________", space_after=2)
    add_p(doc, "Solicitante / Responsable técnico", bold=True, size=10, space_after=2)
    add_p(doc, "Proyecto Suertu2s", size=10, color=COLOR_MUTED, space_after=16)
    add_p(doc, "_______________________________", space_after=2)
    add_p(doc, "Autoriza / Entrega accesos", bold=True, size=10, space_after=2)
    add_p(doc, "Cliente / TI / Titular dominio", size=10, color=COLOR_MUTED, space_after=8)
    add_p(doc, "Fecha: ____ / ____ / ________", size=10, color=COLOR_MUTED)

    doc.add_page_break()

    # 14
    add_heading(doc, "14. Entregables", 1)
    add_plain(
        doc,
        "“Entregables” = lo que se entrega concreto al terminar (o al presentar) el "
        "proyecto: no solo una idea, sino archivos, sistema y documentación.",
    )
    add_bullets(
        doc,
        [
            "Código fuente en GitHub (https://github.com/MathiasAlejandr0/lasuerte).",
            "Aplicación funcional (landing con widget/FAB, checkout, admin con Sorteos, afiliados).",
            "Repositorio dual v1/ (original) y v2/ (versión actual) en GitHub.",
            "Migración SQL única: supabase/migrations/20260811000000_init.sql.",
            "Documentación de presentación y este informe formal.",
            "Pipeline CI (.github/workflows/ci.yml).",
            "Checklist operativo en panel admin.",
        ],
    )

    # 15
    add_heading(doc, "15. Conclusiones y recomendaciones", 1)
    add_p(
        doc,
        "El proyecto Suertu2s constituye una solución informática completa y profesional "
        "para un negocio de packs + sorteo, superando las limitaciones típicas de "
        "WordPress/WooCommerce en control de negocio, seguridad y operación. El desarrollo "
        "está avanzado; el factor crítico restante es la habilitación de infraestructura "
        "y credenciales por parte del cliente/TI.",
        align="justify",
    )
    add_plain(
        doc,
        "Resumen para el equipo completo: la web ya está construida como producto serio, "
        "con landing comercial, consulta de pedidos en la home, botón de soporte, panel de "
        "ciclos de sorteo y comparación v1/v2. Para abrirla al público con cobros reales "
        "necesitamos que nos den las llaves (pagos, base de datos, correo, dominio). "
        "Después: probar, publicar y operar el sorteo.",
    )
    add_p(doc, "Se recomienda:", bold=True, space_after=6)
    add_numbered(
        doc,
        [
            "Aprobar y entregar los accesos de la sección 13 a la brevedad.",
            "Aplicar la migración SQL en un proyecto Supabase limpio.",
            "Configurar webhooks y probar un pago real controlado antes del lanzamiento masivo.",
            "Validar bases legales con asesoría jurídica.",
            "Planificar fase 2: catálogo 100% en Supabase, datos de comprador más completos y 2FA admin.",
        ],
    )

    # 16
    add_heading(doc, "16. Anexos", 1)
    add_heading(doc, "Anexo A — Glosario para todo el staff", 2)
    add_p(
        doc,
        "Lista de términos que aparecen en el proyecto, explicados sin asumir conocimientos técnicos.",
        align="justify",
        italic=True,
        size=10.5,
        color=COLOR_MUTED,
    )
    glossary = [
        ("API", "Puerta de comunicación entre sistemas. Ejemplo: la web habla con Mercado Pago por una API."),
        ("Backend", "La parte del sistema que corre en el servidor (decide precios, guarda pedidos, valida pagos)."),
        ("Frontend", "Lo que ves en pantalla (botones, textos, formularios)."),
        ("Base de datos", "Archivo organizado donde se guardan pedidos, números, afiliados, etc."),
        ("Supabase", "Servicio en la nube que nos da base de datos Postgres lista para usar."),
        ("Postgres / PostgreSQL", "Tipo de base de datos profesional (como un Excel muy potente y seguro)."),
        ("RLS (Row Level Security)", "Reglas de la base de datos que limitan quién puede leer/escribir cada fila."),
        ("Hosting", "Servicio donde “vive” la web en internet."),
        ("Dominio", "El nombre de la web (ej. suertu2s.cl)."),
        ("DNS", "La libreta de direcciones que conecta el dominio con el servidor correcto."),
        ("SSL / HTTPS", "Candadito de seguridad: la comunicación viaja cifrada."),
        ("Checkout", "Pantalla final donde la persona confirma y paga."),
        ("Pasarela de pago", "Empresa que procesa el cobro (Mercado Pago, Webpay/Transbank)."),
        ("Webhook", "Aviso automático de la pasarela al servidor: “este pago quedó aprobado”."),
        ("Mock / pago de prueba", "Simulación de pago para probar en desarrollo. En producción está apagado."),
        ("Pack", "Producto digital (ilustración) que incluye N números de sorteo."),
        ("Ticket / número", "Número de participación asignado al comprador después del pago."),
        ("Fulfill / fulfillment", "Paso en que el pedido se marca pagado y se emiten los números."),
        ("Email de confirmación", "Correo automático con los números e ilustraciones."),
        ("Flag de email", "Marca interna que indica si el correo ya se envió (sirve para reintentar)."),
        ("Admin / panel", "Zona privada (/admin) para operar el negocio."),
        ("Afiliado / referido", "Persona con código propio que trae compradores a cambio de comisión."),
        ("Liquidación / payout", "Registro de comisión pagada a un afiliado."),
        ("Countdown", "Contador regresivo hasta la fecha del sorteo."),
        ("Live / stream", "Transmisión en vivo del sorteo (YouTube/Twitch) embebida en la web."),
        ("Go-live", "Momento en que la web queda pública con dominio y cobros reales."),
        ("Variable de entorno (.env)", "Casilla secreta de configuración (claves) fuera del código público."),
        ("Git / GitHub", "Sistema y plataforma para guardar y versionar el código del proyecto."),
        ("Commit / push", "Guardar un cambio en el historial del código y subirlo al repositorio."),
        ("CI (Integración Continua)", "Pruebas automáticas que corren al subir cambios (para detectar errores)."),
        ("Middleware", "Filtro de seguridad que revisa pedidos a la web antes de llegar a la función final."),
        ("CSRF", "Ataque donde un sitio malicioso intenta hacer acciones en tu nombre; aquí se mitiga."),
        ("HMAC / firma", "Sello criptográfico para verificar que un mensaje no fue alterado."),
        ("Cookie de sesión", "Dato que el navegador guarda para mantenerte logueado de forma segura."),
        ("httpOnly", "La cookie no la puede leer JavaScript del navegador (más seguro)."),
        ("Hash / scrypt", "Forma de guardar contraseñas sin guardar la contraseña en texto claro."),
        ("CSP / HSTS", "Cabeceras de seguridad del navegador que reducen ciertos ataques."),
        ("Rate limit", "Límite de intentos (ej. no permitir mil logins por minuto desde la misma IP)."),
        ("TypeScript", "Lenguaje que ayuda a detectar errores de programación antes de publicar."),
        ("Next.js", "Framework (base) con el que está hecha esta web."),
        ("Vercel", "Ejemplo de plataforma donde se puede publicar una app Next.js."),
        ("Resend", "Servicio usado para enviar correos transaccionales."),
        ("Migración SQL", "Script que crea/actualiza las tablas de la base de datos."),
        ("Seed", "Datos iniciales de ejemplo (packs, afiliados demo, etc.)."),
        ("KPI", "Indicador de negocio (ej. ingresos, pedidos pagados, conversión)."),
        ("Analítica", "Pantallas/números para entender cómo va la venta y el sorteo."),
        ("Stakeholder", "Persona interesada o responsable del proyecto (dueño, TI, operación)."),
    ]
    for term, meaning in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(f"{term}: ")
        set_font(r1, FONT_TITLE, bold=True, size=10.5, color=COLOR_GREEN)
        r2 = p.add_run(meaning)
        set_font(r2, FONT_BODY, size=10.5, color=COLOR_TEXT)

    add_heading(doc, "Anexo B — Referencias del repositorio", 2)
    add_bullets(
        doc,
        [
            "Código: https://github.com/MathiasAlejandr0/lasuerte",
            "Presentación corta: docs/PRESENTACION-PROYECTO.md",
            "Esquema DB: supabase/migrations/20260811000000_init.sql",
            "Ejemplo de entorno: .env.example",
        ],
    )
    add_heading(doc, "Anexo C — Declaración", 2)
    add_p(
        doc,
        "Este documento describe el estado del proyecto Suertu2s como producto "
        "informático serio: alcance funcional, arquitectura, costos referenciales de "
        "mercado en Chile, riesgos, seguridad y solicitud formal de accesos. Los montos "
        "de desarrollo son referenciales de mercado y no constituyen cotización vinculante "
        "salvo acuerdo escrito posterior. Las explicaciones en lenguaje simple buscan que "
        "todo el equipo —no solo informática— pueda tomar decisiones informadas.",
        align="justify",
    )

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
